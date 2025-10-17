import io
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

def _p(doc: Document, text: str, bold=False, align=None, space_after=4):
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.bold = bold
    para.paragraph_format.space_after = Pt(space_after)
    if align:
        para.alignment = align
    return para

def build_docx_from_body(ciudad: str, fecha: str, destinatario: str, paciente: str, hc: str, body: str) -> bytes:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    _p(doc, f"{ciudad}, {fecha}", align=WD_ALIGN_PARAGRAPH.RIGHT)
    _p(doc, "Señor", bold=False)
    _p(doc, destinatario or "—", bold=False)
    _p(doc, "Presente.-", bold=False)
    _p(doc, f"H.C.Nº {hc or '—'}", bold=False)
    doc.add_paragraph("")

    _p(doc, f"Estimado señor {paciente.split()[-1] if paciente else ''}:", bold=False)

    intro = ("A continuación, detallamos el resultado de su examen médico preventivo y ocupacional. ")
    _p(doc, intro, bold=False)

    # Cuerpo generado por LLM
    for line in body.split("\n"):
        if line.strip():
            _p(doc, line.strip())

    _p(doc, "Si deseara alguna aclaración o información adicional con respecto al presente informe, no dude en llamarnos. "
             "Los médicos de Doktuz estamos a su disposición.", bold=False)

    _p(doc, "Atentamente,", bold=False)
    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()

# Optional template fill: replace simple placeholders in a .docx
def fill_template_docx(template_path: str, ciudad: str, fecha: str, destinatario: str, paciente: str, hc: str, body: str) -> bytes:
    doc = Document(template_path)
    placeholders = {
        "{{CIUDAD}}": ciudad,
        "{{FECHA}}": fecha,
        "{{DESTINATARIO}}": destinatario or "—",
        "{{PACIENTE}}": paciente or "—",
        "{{HC}}": hc or "—",
        "{{CUERPO}}": body,
    }
    for p in doc.paragraphs:
        for k,v in placeholders.items():
            if k in p.text:
                inline = p.runs
                for i in range(len(inline)):
                    if k in inline[i].text:
                        inline[i].text = inline[i].text.replace(k, v)

    # tables too
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for k,v in placeholders.items():
                    if k in cell.text:
                        for r in cell.paragraphs[0].runs:
                            r.text = r.text.replace(k, v)
    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()
