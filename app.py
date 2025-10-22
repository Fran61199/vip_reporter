#!/usr/bin/env python3
# -*- coding: utf-8 -*-

import os, re, io, zipfile, traceback, json
from datetime import datetime
from flask import Flask, request, render_template, send_file, jsonify
import pdfplumber
from docx import Document

# === Rutas absolutas para encontrar la plantilla ===
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
TEMPLATE_FOLDER_USER = os.getenv("TEMPLATE_FOLDER_USER", "user_templates")
USER_TEMPLATES_DIR = os.path.join(BASE_DIR, TEMPLATE_FOLDER_USER)

# Candidatos de plantilla (en orden de prioridad)
CANDIDATE_PLANTILLAS = [
    os.path.join(USER_TEMPLATES_DIR, "plantilla.docx"),
    os.path.join(BASE_DIR, "plantilla.docx"),
    os.path.join(BASE_DIR, "templates", "plantilla.docx"),
    os.path.join(BASE_DIR, "template", "plantilla.docx"),
]
PLANTILLA_PATH = next((p for p in CANDIDATE_PLANTILLAS if os.path.exists(p)), None)

# ====== Config ======
MODEL_NAME = os.getenv("OPENAI_MODEL", "gpt-4o-mini")
MAX_CHARS_GPT = 12000
ESTIMAR_MODO = "limpio"  # "limpio" | "claves"
MAX_PDFS_IA = int(os.getenv("MAX_PDFS_IA", "13"))  # ← umbral para usar IA

# ====== App / carpetas ======
app = Flask(__name__)
UPLOAD_FOLDER = "uploads"
OUTPUT_FOLDER = "outputs"
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(OUTPUT_FOLDER, exist_ok=True)

# ====== Limpieza de texto ======
PATTERNS = [
    r'https?://\S+',
    r'===\s*Página.*?===',
    r'\d{2}/\d{2}/\d{2},\s*\d{1,2}:\d{2}\s*(a|p)\.m\.',
    r'\(\*\)\s*Dato\s*Modificado',
    r'Firma y Sello.*',
    r'Nombre,?\s*Firma.*',
    r'Formato de Impresion',
    r'Fecha de (Registro|Revisión).*',
    r'\s{2,}'
]
def limpiar_texto(texto: str) -> str:
    for p in PATTERNS:
        texto = re.sub(p, '', texto, flags=re.IGNORECASE)
    texto = texto.replace(' \n', '\n').replace('\n ', '\n')
    texto = re.sub(r'[ \t]+', ' ', texto)
    texto = re.sub(r'(?i)\b(CONCLUSIONES|RECOMENDACIONES|DIAGN[ÓO]STICOS?)\b', r'\n\n\1\n', texto)
    texto = re.sub(r'\n{3,}', '\n\n', texto)
    return texto.strip()

# ====== Clasificación ======
CATEGORIAS = [
    ("LABORATORIO", [r'INFORME DE LABORATORIO', r'EXAMEN DE LABORATORIO', r'HEMOGRAMA', r'BIOQUIMICA']),
    ("CARDIOLOGIA", [r'ELECTROCARDIOGRAF', r'CARDIOVASCULAR', r'EKG']),
    ("AUDIOLOGIA", [r'AUDIOL', r'OTOSCOPIA', r'OTORRINO']),
    ("DERMATOLOGIA", [r'DERMATOL']),
    ("ECOGRAFIA ABDOMINAL", [r'ECOGRAF(IA|ICA)\s+ABDOM']),
    ("ECOGRAFIA PELVICA", [r'ECOGRAF(IA|ICA)\s+PELV']),
    ("HISTORIA CLINICA", [r'HISTORIA CLINICA MEDICA OCUPACIONAL', r'HISTORIA CL[IÍ]NICA']),
    ("MUSCULO ESQUELETICO", [r'MUSCULO', r'ESQUEL[EÉ]TICA']),
    ("NEUROLOGIA", [r'NEUROL']),
    ("ODONTOLOGIA", [r'ODONTO', r'ODONTOGRAMA']),
    ("OFTALMOLOGIA", [r'OFTALMO']),
    ("PRUEBA DE ESFUERZO", [r'PRUEBA DE ESFUERZO', r'PROTOCOLO BRUCE']),
    ("PSICOLOGIA", [r'PSICOL', r'EPWORTH']),
    ("UROLOGIA", [r'UROL']),
    ("RADIOLOGIA", [r'RADIOGRA', r'TÓRAX']),
    ("ESPIROMETRIA", [r'ESPIROM']),
]
def clasificar(texto: str) -> str:
    upper = texto.upper()
    for nombre, pats in CATEGORIAS:
        for pat in pats:
            if re.search(pat, upper):
                return nombre
    return "OTROS"

# ====== Extracción PDF ======
def extraer_texto_pdf(pdf_path: str) -> str:
    partes = []
    with pdfplumber.open(pdf_path) as pdf:
        for page in pdf.pages:
            txt = page.extract_text(x_tolerance=2, y_tolerance=2) or ""
            txt = txt.strip()
            if txt:
                partes.append(txt)
    return "\n\n".join(partes)

# ====== Estimación (solo UI) ======
def estimar_tokens_aprox(texto: str) -> int:
    return max(1, len(texto) // 4)  # ~1 token ≈ 4 chars

def extraer_solo_claves(texto: str) -> str:
    keep = []
    for tag in [r'(?is)\bCONCLUSIONES\b.*?(?=\n[A-ZÁÉÍÓÚÑ ]{3,}|$)',
                r'(?is)\bDIAGN[ÓO]STICOS?\b.*?(?=\n[A-ZÁÉÍÓÚÑ ]{3,}|$)',
                r'(?is)\bRECOMENDACIONES\b.*?(?=\n[A-ZÁÉÍÓÚÑ ]{3,}|$)']:
        m = re.search(tag, texto)
        if m:
            keep.append(m.group(0))
    if not keep:
        keep = ["\n".join(texto.splitlines()[:8])]
    s = "\n\n".join(keep)
    s = re.sub(r'\s{2,}', ' ', s)
    return s.strip()

# ====== DOCX ======
def _replace_in_paragraph(paragraph, mapping):
    for key, val in mapping.items():
        placeholder = f'{{{{{key}}}}}'
        if placeholder in paragraph.text:
            paragraph.text = paragraph.text.replace(placeholder, val)

def _replace_in_document(doc: Document, mapping: dict):
    for p in doc.paragraphs:
        _replace_in_paragraph(p, mapping)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    _replace_in_paragraph(p, mapping)

def generar_docx(datos: dict, plantilla_path: str, salida_path: str):
    from docx.shared import Pt
    doc = Document(plantilla_path)
    _replace_in_document(doc, datos)
    # Forzar Calibri 11
    for p in doc.paragraphs:
        for run in p.runs:
            run.font.name = "Calibri"
            run.font.size = Pt(11)
    for table in doc.tables:
        for row in table.rows:
            for cell in table.rows[0].table.rows[0].cells if False else []:
                pass
        for row in doc.tables:
            for r in row.rows:
                for cell in r.cells:
                    for p in cell.paragraphs:
                        for run in p.runs:
                            run.font.name = "Calibri"
                            run.font.size = Pt(11)
    doc.save(salida_path)

# ====== GPT ======
def redactar_vip_con_gpt(paciente: str, insumo: str) -> str:
    api_key = os.getenv("OPENAI_API_KEY") or os.getenv("openai_api_key")
    if not api_key:
        raise RuntimeError("OPENAI_API_KEY no configurada")
    try:
        from openai import OpenAI
        client = OpenAI(
            api_key=api_key,
            timeout=float(os.getenv("OPENAI_REQUEST_TIMEOUT", "60"))  # opcional
        )
        insumo = insumo[:MAX_CHARS_GPT]

        system_msg = (
  "Eres médico ocupacional y redactas informes VIP con tono profesional, claro y empático. "
  "Escribe SIEMPRE en SEGUNDA PERSONA de cortesía (usted). "
  "No uses 'el paciente'/'la paciente' ni tercera persona. "
  "Estructura en párrafos independientes, sin títulos gruesos, salvo la sección final de recomendaciones en viñetas. "
  "Incluye valores y unidades cuando estén en el insumo y ofrece interpretaciones breves. "
  "No inventes nada: si un dato no aparece, no lo supongas ni lo rellenes. "
  "Respeta la evidencia del insumo y evita repetir texto administrativo. "
  "ESPAÑOL neutro de salud ocupacional."
)
        prompt = f"""Redacta el CUERPO de un informe VIP con el siguiente insumo clínico ya limpio.
Tono y forma:
- Segunda persona (usted) SIEMPRE. Evita frases como “el paciente…”.
- Párrafos breves (3–5 oraciones), separados por líneas en blanco.
- Cierre con recomendaciones en viñetas (máximo 5), claras y accionables.

Orden sugerido:
1) Contexto general: edad y antecedentes relevantes (familiares, personales, alergias, cirugías si aplica).
2) Examen físico: peso/talla/IMC si están disponibles, PA, FC; interpretación breve.
3) Hallazgos por especialidad (solo si aparecen en el insumo), por ejemplo:
   - Oftalmología
   - Cardiología (EKG/prueba de esfuerzo)
   - Audiología/ORL
   - Tórax/Espirometría
   - Radiología (incluye columna si procede)
   - Odontología
   - Músculo-esquelético
   - Ecografías (abdominal/pélvica)
   - Urología
4) Laboratorio: hemograma, glucosa, perfil lipídico, orina, PSA/tiroides/marcadores (solo lo disponible).
5) Cierre + “En conclusión” y recomendaciones en viñetas.

Reglas:
- No repitas nombres de empresa, sellos ni datos administrativos.
- Usa cifras con unidades y corta interpretación (p.ej., “IMC 26.2 kg/m²: rango de sobrepeso”).
- Si una especialidad no aparece en el insumo, omítela.
- Evita listados largos dentro de párrafos; prioriza claridad clínica.
- No inventes rangos de referencia si no están en el insumo.

INSUMO:
{insumo}
""".strip()

        resp = client.chat.completions.create(
            model=MODEL_NAME,
            temperature=0.4,
            max_tokens=900,
            messages=[
                {"role": "system", "content": system_msg},
                {"role": "user", "content": prompt}
            ]
        )

        texto = resp.choices[0].message.content.strip()
        partes = texto.split("\n")
        limpio = []
        en_reco = False
        for ln in partes:
            if ln.strip().lower().startswith("recomendaciones"):
                en_reco = True
                limpio.append("Recomendaciones:")
                continue
            if not en_reco and (ln.strip().startswith("•") or ln.strip().startswith("- ")):
                ln = ln.lstrip("•- ").strip()
            limpio.append(ln)
        return "\n".join(limpio).strip()
    except Exception as e:
        print("Error GPT:", e)
        traceback.print_exc()
        raise

# ========= Rutas =========

@app.route("/", methods=["GET"])
def home():
    return render_template("index.html")

@app.route("/generate-docx", methods=["POST"])
def generate_docx():
    # Campos
    nombre = request.form.get("nombre", "").strip()
    ape1   = request.form.get("ape1", "").strip()
    ape2   = request.form.get("ape2", "").strip()
    hc     = request.form.get("hc", "").strip()
    sexo   = request.form.get("sexo", "M").strip().upper()

    if not nombre or not ape1 or not hc:
        return "Faltan campos obligatorios (Nombre, Primer Apellido, H.C.).", 400

    archivos = request.files.getlist("pdfs")
    if not archivos or all(not f.filename.lower().endswith(".pdf") for f in archivos):
        return "Debes adjuntar al menos un PDF válido.", 400

    trato = "Señor" if sexo == "M" else "Señora"
    saludo_adj = "Estimado" if sexo == "M" else "Estimada"
    apellido = (ape1.split()[0] if ape1 else "").capitalize()

    # Fecha Lima
    meses = ["enero","febrero","marzo","abril","mayo","junio","julio","agosto","septiembre","octubre","noviembre","diciembre"]
    hoy = datetime.now()
    fecha_str = f"Lima, {hoy.day} de {meses[hoy.month-1]} del {hoy.year}"

    # Procesar PDFs y armar insumo
    cat_a_texto = {}
    total_bytes = 0
    for i, pdf in enumerate(archivos):
        if not pdf.filename.lower().endswith(".pdf"):
            continue
        ruta_pdf = os.path.join(UPLOAD_FOLDER, pdf.filename)
        pdf.save(ruta_pdf)
        try:
            total_bytes += os.path.getsize(ruta_pdf)
        except OSError:
            pass

        crudo = extraer_texto_pdf(ruta_pdf)
        limpio = limpiar_texto(crudo)
        categoria = clasificar(limpio)

        salida_fn = f"{i+1:02d}_{categoria}.txt"
        ruta_txt = os.path.join(OUTPUT_FOLDER, salida_fn)
        with open(ruta_txt, "w", encoding="utf-8") as f:
            f.write(limpio)

        cat_a_texto.setdefault(categoria, []).append(limpio)

    bloques = []
    for cat, textos in cat_a_texto.items():
        bloques.append(f"### {cat}\n" + "\n".join(textos))
    insumo_completo = "\n\n".join(bloques).strip()

    if not insumo_completo:
        return "Error: no se pudo construir el texto base de los PDFs.", 400

    paciente = f"{nombre} {ape1} {ape2}".strip()

    # ----- CLAVE PARA DESCARTAR RAM/TIMEOUT -----
    def _resumen_local(catmap):
        partes = []
        for cat, textos in catmap.items():
            joined = "\n".join(textos)
            snippet = "\n".join(joined.splitlines()[:3]).strip()
            snippet = re.sub(r'\s+', ' ', snippet)[:500]
            partes.append(f"{cat}: {snippet}")
        return "\n\n".join(partes) or "(Contenido pendiente)"

    if len(archivos) > MAX_PDFS_IA:
        # Carga alta → NO usar IA (evita 500 por timeout/RAM)
        cuerpo = "(Se omitió IA por carga alta)\n\n" + _resumen_local(cat_a_texto)
    else:
        if not os.getenv("OPENAI_API_KEY") and not os.getenv("openai_api_key"):
            return "Error: OPENAI_API_KEY no configurada en el entorno.", 400
        cuerpo = redactar_vip_con_gpt(paciente, insumo_completo)
    # ---------------------------------------------

    # Generar DOCX
    datos = {
        "FECHA": fecha_str,
        "SEXO_TRATO": trato,
        "NOMBRE_COMPLETO": paciente,
        "HC": hc,
        "SEXO_ADJETIVO": saludo_adj,
        "APELLIDO": apellido,
        "CUERPO": cuerpo or "(Contenido pendiente)"
    }

    total_mb = round(total_bytes / (1024*1024), 2)
    nombre_docx = f"Informe_{apellido or 'Paciente'}_{total_mb}MB.docx"
    salida_docx = os.path.join(OUTPUT_FOLDER, nombre_docx)

    if not PLANTILLA_PATH or not os.path.exists(PLANTILLA_PATH):
        return "Error: no se encontró 'plantilla.docx'. Ubícala en user_templates/, templates/ o raíz.", 500
    generar_docx(datos, PLANTILLA_PATH, salida_docx)

    # Sanity check antes de enviar
    try:
        sz = os.path.getsize(salida_docx)
        if sz <= 0:
            return "El DOCX se generó vacío o no existe.", 500
    except Exception as e:
        return f"No se pudo acceder al DOCX generado: {e}", 500

    resp = send_file(salida_docx, as_attachment=True, download_name=nombre_docx)
    resp.headers["X-Upload-MB"] = str(total_mb)
    resp.headers["X-PDF-Count"] = str(len(archivos))
    resp.headers["X-IA-Used"] = "false" if len(archivos) > MAX_PDFS_IA else "true"
    return resp

if __name__ == "__main__":
    app.run(debug=True)
